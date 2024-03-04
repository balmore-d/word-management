from typing import Union, Dict
from fastapi import FastAPI, Form, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx2pdf import convert


class FileConversion(BaseModel):
    base64: str | None = None


class ReplacementMap(BaseModel):
    replacements: Dict[str, str]


app = FastAPI()

origins = [
    "http://localhost:8080/*",
    "http://127.0.0.1:8080/*"
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


@app.get("/")
def read_root():
    return {"Hello": "World"}


@app.get("/item/{item_id}")
def read_item(item_id: int, q: Union[str, None] = None):
    return {"item_id": item_id, "q": q}


def make_replace(runs, open_index, close_index, value_map):
    print("yes")
    op_index_txt = runs.get(open_index)
    pass


def update_runs(runs, value_map):
    print("	Entered updating runs meth")
    var_to_replace = ""
    is_reading_var = False
    var_open_run_num = 0
    var_close_run_num = 0
    is_first_op_brace_inline = False
    is_first_close_brace_inline = False
    print("total run inside paragraph: ", len(runs))
    run_count = 0
    for run in runs:
        text = run.text
        try:
            if text != "" and text is not None:
                if "{{" in text and "}" not in text and "}}" not in text:
                    print("GENERALITÈS:", text, text.find("{{"), len(text))
                    if text.find("{{") == 0 and len(text) > 2:
                        print("{{xx", "/", text)
                    else:
                        print("xxx {{ || {{", "/", text)
                #TODO: do rest
                elif "{{" in text and "}" in text and "}}" not in text:
                    print("GENERALITÈS:", text, text.find("{{"), text.find("}"), len(text))
                    if text.find("{{") < text.find("}"):
                        print('{{xxx}', "/", text)
                    else:
                        print("} xxx {{", "/", text)
                elif "{{" in text and "}}" in text:
                    print("GENERALITÈS:", text, text.find("{{"), text.find("}}"), len(text))
                    if text.find("{{") < text.find("}}"):
                        print("{{xxx}}", "|", text)
                        for key, value in value_map.items():
                            if key in text:
                                sub_str_b4_op_braces = ""
                                sub_str_after_close_braces = ""
                                run_text_length = len(text)
                                opening_index = text.find("{{")
                                closing_index = text.find("}}")
                                if opening_index > 0:
                                    sub_str_b4_op_braces = text[:opening_index]
                                if run_text_length - closing_index > 2:
                                    sub_str_after_close_braces = text[closing_index + 2:]
                                text = ''.join([sub_str_b4_op_braces, value_map.get(key), sub_str_after_close_braces])
                                run.text = text
                            #end if
                        #end for
                        if "{" in text or "{{" in text:
                            print(
                                "warning! one complete variable was read and replaced, but another began in same run")
                            is_first_op_brace_inline = "{" in text and "{{" not in text
                            is_reading_var = True
                    else:
                        print("}} xxx {{", "/", text)
                elif "}}" in text and "{" not in text and "{{" not in text:
                    if text.find("}}") == 0 and len(text) > 2:
                        print("}} xxx", "/", text)
                    else:
                        print("xxx}}", "/", text)
                elif "{" in text and "}}" in text:
                    if text.find("{") < text.find("}}"):
                        print("{xxx}}", "/", text)
                    else:
                        print("}} xxx {", "/", text)
                elif "{" in text and "}" in text and "}}" not in text:
                    if text.find("{") == 0 and len(text) > 1:
                        print("{text", "/", text)
                    else:
                        print("xxx { || {", "/", text)
                elif "{" in text and "}" in text and "}}" not in text:
                    if text.find("}") < text.find("{"):
                        print("} xxx {", "/", text)
                    else:
                        print("{xxx}", "/", text)
                elif "}" in text and "{" not in text and "{{" not in text:
                    if text.find("{") == 0:
                        print("} xxx || }", "/", text)
                    else:
                        print("text}", "/", text)
                else:
                    # !!! this will print IF reading var flag is True
                    print("bare text ... reading var", "/", text)
            else:
                print("the run text was: ", "/", text)
        except Exception as e:
            print("text with glitch: ", text, e)


def replace_variables_in_docx(docx_file, replacements):
    print("entered variables replacement method")
    doc = Document(docx_file)
    #Replace in paragraphs
    for paragraph in doc.paragraphs:
        update_runs(paragraph.runs, replacements)

    #Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            new_text = paragraph.text.replace(key, value)
                            for run in paragraph.runs:
                                run.text = new_text
                                break
    return doc


@app.post("/convertDocxBase64ToPdf")
async def convert_docx_to_pdf(file: UploadFile):
    try:
        print("docxBase64 received ", file.filename)
        with open("temp.docx", "wb") as f:
            f.write(await file.read())
        replacements = {
            "nombres_doc_cliente": "Balmore",
            "apellidos_doc_cliente": "Ortíz",
            "texto_conocido_por": "balmore_",
            "codigo_cliente": "OM17007",
            "no_inventario": "12345"
        }
        doc = replace_variables_in_docx("temp.docx", replacements)
        doc.save("modified.docx")
        convert("modified.docx")
        return {"filename": file.filename, "message": "converted successfully"}
    except Exception as ex:
        print("docxBase64 could not be converted: ", ex)
        return {"error": str(ex)}
    finally:
        import os
        os.remove("temp.docx")
        os.remove("modified.docx")
